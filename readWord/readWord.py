import json

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import pprint as pp
"""
Generate a reference to each paragraph and table child within *parent*,
in document order. Each returned value is an instance of either Table or
Paragraph. *parent* would most commonly be a reference to a main
Document object, but also works for a _Cell object, which itself can
contain paragraphs and tables.
"""
def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def convertDocxToText(path):
    doc = Document(path)
    fullText = []


    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            if block.text.find('Date') != -1:
                start_idx = block.text.find(':')
                print(block.text[start_idx + 1:].strip())

            if block.text.find('Ref') != -1:
                start_idx = block.text.find(':')
                print(block.text[start_idx + 1:].strip())
            # print(f"block.text : {block.text}")

        elif isinstance(block, Table):
            for i, row in enumerate(block.rows):
                text = (cell.text for cell in row.cells)
                # .... This part is for Table.

                if i == 0:
                    keys = tuple(text)
                    continue

                row_data = dict(zip(keys, text))
                for k,v in row_data.items():
                    print(f"{k} : {v}")
                print("-------------------")
                # pp.pprint(row_data)
                # print(str(row_data))
                fullText.append(str(row_data))

    return '\n'.join(fullText)

FILE_NAME = 'CN0810-DV-R0.docx'
# FILE_NAME = 'LN0725-DV(A)-R0.docx'
# FILE_NAME = 'LN0819-CL(A)-R0.docx'
# FILE_NAME = 'TK0804-DV(A)-R0.docx'
# FILE_NAME = 'LN1333-DV(A)-R0.docx'
# doc = Document(FILE_NAME)
# convertDocxToText(FILE_NAME)

with open('CL_DV_categorization.json') as f:
    CL_DV_category_dict = json.loads(f.read())

pp.pprint(CL_DV_category_dict)

"""
for block in iter_block_items(doc):

    if isinstance(block, Paragraph):
        print(block.text)
    elif isinstance(block, Table):
        for row in block.rows:
            row_data = []
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    row_data.append(paragraph.text)
            print((row_data))
"""


"""
import pandas as pd
table = doc.tables[0]

data = []

keys = None
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)

    if i == 0:
        keys = tuple(text)
        continue
    row_data = dict(zip(keys, text))
    data.append(row_data)


df = pd.DataFrame(data)

print(df.iloc[0,:].tolist()) #row index 0
print(df.iloc[:,0].tolist()) #column index 0
print(df.iloc[0,0]) #cell(0,0)
print(df.iloc[1:,2].tolist()) #column index 2, but skip first row
"""

